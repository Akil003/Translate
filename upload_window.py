import math
from random import random
import speech_recognition as sr
import login as lg
import os
from os import path
import pydub
from playsound import playsound
from PyQt5 import QtCore, QtGui, QtWidgets
from Libraries.image_to_text import *
from googletrans import Translator
from Libraries.Translate import *
from docx import Document
import PyPDF2
from random import randint

import Libraries.userdetails as userdetails


translator = Translator()


class Ui_UploadWindow(object):

    # 1 for go back
    # 2 for log out
    def open_window(self, val):
        self.login_screen = QtWidgets.QMainWindow()
        self.login_screen.setWindowTitle("Translate+")
        self.ui = lg.Ui_login_window()
        self.ui.setupUi(self.login_screen)

        if val == 1:
            self.ui.mode_frame.setVisible(True)
        elif val == 2:
            self.ui.mode_frame.setVisible(False)

        self.login_screen.show()

    def setupUi(self, UploadWindow):
        self.lscr = UploadWindow

        UploadWindow.setObjectName("UploadWindow")
        UploadWindow.resize(813, 600)
        UploadWindow.setStyleSheet("background-color: rgb(255, 255, 255);\n"
                                   "\n"
                                   "QPushButton {\n"
                                   "    font: 12pt \"Comic Sans MS\";\n"
                                   "color : white;\n"
                                   "background-color: #4CAF50;\n"
                                   "border-radius: 2px;\n"
                                   "border: 2px solid #4CAF50;\n"
                                   "}\n"
                                   "\n"
                                   "QPushButton:hover{\n"
                                   "    background-color: white;\n"
                                   "    color: black;\n"
                                   "}\n"
                                   "")
        self.centralwidget = QtWidgets.QWidget(UploadWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.main_frame = QtWidgets.QFrame(self.centralwidget)
        self.main_frame.setGeometry(QtCore.QRect(210, 100, 371, 411))
        self.main_frame.setStyleSheet("QFrame#main_frame {\n"
                                      "    border: 3px solid #4CAF50;\n"
                                      "}")
        self.main_frame.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.main_frame.setFrameShadow(QtWidgets.QFrame.Raised)
        self.main_frame.setObjectName("main_frame")
        self.recording_label = QtWidgets.QLabel(self.main_frame)
        self.recording_label.setGeometry(QtCore.QRect(20, 30, 340, 91))
        font = QtGui.QFont()
        font.setFamily("Comic Sans MS")
        font.setPointSize(12)
        self.recording_label.setFont(font)
        self.recording_label.setObjectName("recording_label")
        self.text_bton = QtWidgets.QPushButton(self.main_frame)
        self.text_bton.setGeometry(QtCore.QRect(110, 140, 160, 50))
        font = QtGui.QFont()
        font.setFamily("Comic Sans MS")
        font.setPointSize(12)
        font.setBold(False)
        font.setItalic(False)
        font.setWeight(50)
        self.text_bton.setFont(font)
        self.text_bton.setStyleSheet("QPushButton {\n"
                                     "    font: 12pt \"Comic Sans MS\";\n"
                                     "color : white;\n"
                                     "background-color: #4CAF50;\n"
                                     "border-radius: 2px;\n"
                                     "border: 2px solid #4CAF50;\n"
                                     "}\n"
                                     "\n"
                                     "QPushButton:hover{\n"
                                     "    background-color: white;\n"
                                     "    color: black;\n"
                                     "}\n"
                                     "")
        self.text_bton.setIconSize(QtCore.QSize(100, 100))
        self.text_bton.setFlat(False)
        self.text_bton.setObjectName("text_bton")
        self.image_bton = QtWidgets.QPushButton(self.main_frame)
        self.image_bton.setGeometry(QtCore.QRect(110, 300, 160, 50))
        font = QtGui.QFont()
        font.setFamily("Comic Sans MS")
        font.setPointSize(12)
        font.setBold(False)
        font.setItalic(False)
        font.setWeight(50)
        self.image_bton.setFont(font)
        self.image_bton.setStyleSheet("QPushButton {\n"
                                      "    font: 12pt \"Comic Sans MS\";\n"
                                      "color : white;\n"
                                      "background-color: #4CAF50;\n"
                                      "border-radius: 2px;\n"
                                      "border: 2px solid #4CAF50;\n"
                                      "}\n"
                                      "\n"
                                      "QPushButton:hover{\n"
                                      "    background-color: white;\n"
                                      "    color: black;\n"
                                      "}\n"
                                      "")
        self.image_bton.setIconSize(QtCore.QSize(100, 100))
        self.image_bton.setFlat(False)
        self.image_bton.setObjectName("image_bton")
        self.audio_bton = QtWidgets.QPushButton(self.main_frame)
        self.audio_bton.setGeometry(QtCore.QRect(110, 220, 160, 50))
        font = QtGui.QFont()
        font.setFamily("Comic Sans MS")
        font.setPointSize(12)
        font.setBold(False)
        font.setItalic(False)
        font.setWeight(50)
        self.audio_bton.setFont(font)
        self.audio_bton.setStyleSheet("QPushButton {\n"
                                      "    font: 12pt \"Comic Sans MS\";\n"
                                      "color : white;\n"
                                      "background-color: #4CAF50;\n"
                                      "border-radius: 2px;\n"
                                      "border: 2px solid #4CAF50;\n"
                                      "}\n"
                                      "\n"
                                      "QPushButton:hover{\n"
                                      "    background-color: white;\n"
                                      "    color: black;\n"
                                      "}\n"
                                      "")
        self.audio_bton.setIconSize(QtCore.QSize(100, 100))
        self.audio_bton.setFlat(False)
        self.audio_bton.setObjectName("audio_bton")
        self.logout_ico = QtWidgets.QLabel(self.centralwidget)
        self.logout_ico.setGeometry(QtCore.QRect(640, 20, 41, 41))
        self.logout_ico.setText("")
        self.logout_ico.setPixmap(QtGui.QPixmap("images/log_out_icon.png"))
        self.logout_ico.setScaledContents(True)
        self.logout_ico.setObjectName("logout_ico")
        self.goback_btn = QtWidgets.QPushButton(self.centralwidget)
        self.goback_btn.setGeometry(QtCore.QRect(320, 520, 160, 50))
        font = QtGui.QFont()
        font.setFamily("Comic Sans MS")
        font.setPointSize(12)
        font.setBold(False)
        font.setItalic(False)
        font.setWeight(50)
        self.goback_btn.setFont(font)
        self.goback_btn.setStyleSheet("QPushButton {\n"
                                      "    font: 12pt \"Comic Sans MS\";\n"
                                      "color : white;\n"
                                      "background-color: #4CAF50;\n"
                                      "border-radius: 2px;\n"
                                      "border: 2px solid #4CAF50;\n"
                                      "}\n"
                                      "\n"
                                      "QPushButton:hover{\n"
                                      "    background-color: white;\n"
                                      "    color: black;\n"
                                      "}\n"
                                      "")
        self.goback_btn.setIconSize(QtCore.QSize(100, 100))
        self.goback_btn.setFlat(False)
        self.goback_btn.setObjectName("goback_btn")
        self.logout_btn = QtWidgets.QPushButton(self.centralwidget)
        self.logout_btn.setGeometry(QtCore.QRect(690, 20, 102, 40))
        font = QtGui.QFont()
        font.setFamily("Comic Sans MS")
        font.setPointSize(12)
        font.setBold(False)
        font.setItalic(False)
        font.setWeight(50)
        self.logout_btn.setFont(font)
        self.logout_btn.setStyleSheet("QPushButton {\n"
                                      "    font: 12pt \"Comic Sans MS\";\n"
                                      "    background-color: white;\n"
                                      "    color: black;\n"
                                      "    border: 2px solid #4CAF50;\n"
                                      "    border-radius: 5%;\n"
                                      "}\n"
                                      "\n"
                                      "QPushButton:hover{\n"
                                      "    background-color: #4CAF50;\n"
                                      "    color: white;\n"
                                      "\n"
                                      "    border: 2px solid #4CAF50;\n"
                                      "}\n"
                                      "")
        self.logout_btn.setIconSize(QtCore.QSize(100, 100))
        self.logout_btn.setFlat(False)
        self.logout_btn.setObjectName("logout_btn")
        self.proj_ico = QtWidgets.QLabel(self.centralwidget)
        self.proj_ico.setGeometry(QtCore.QRect(20, 10, 90, 80))
        self.proj_ico.setText("")
        self.proj_ico.setPixmap(QtGui.QPixmap("images/logo.png"))
        self.proj_ico.setScaledContents(True)
        self.proj_ico.setObjectName("proj_ico")

        self.selection_frame = QtWidgets.QFrame(self.centralwidget)
        self.selection_frame.setGeometry(QtCore.QRect(570, 290, 231, 291))
        self.selection_frame.setStyleSheet("QFrame#selection_frame {\n"
                                           "    border: 3px solid #4CAF50;\n"
                                           "}")
        self.selection_frame.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.selection_frame.setFrameShadow(QtWidgets.QFrame.Raised)
        self.selection_frame.setObjectName("selection_frame")
        self.blank_msg = QtWidgets.QLabel(self.selection_frame)
        # self.blank_msg.setText("Not working")
        # self.blank_msg.setGeometry(QtCore.QRect(10, 60, 210, 30))

        self.sec_lang = QtWidgets.QComboBox(self.selection_frame)
        self.sec_lang.setGeometry(QtCore.QRect(10, 150, 210, 30))
        font = QtGui.QFont()
        font.setFamily("Comic Sans MS")
        font.setPointSize(10)
        self.sec_lang.setFont(font)
        self.sec_lang.setObjectName("to_lang")

        self.primary_lang = QtWidgets.QComboBox(self.selection_frame)
        self.primary_lang.setGeometry(QtCore.QRect(10, 60, 210, 30))
        font = QtGui.QFont()
        font.setFamily("Comic Sans MS")
        font.setPointSize(10)
        self.primary_lang.setFont(font)
        self.primary_lang.setObjectName("from_lang")

        self.to_lbl = QtWidgets.QLabel(self.selection_frame)
        self.to_lbl.setGeometry(QtCore.QRect(20, 110, 40, 31))
        font = QtGui.QFont()
        font.setFamily("Comic Sans MS")
        font.setPointSize(10)
        self.to_lbl.setFont(font)
        self.to_lbl.setObjectName("to_lbl")
        self.from_btn = QtWidgets.QLabel(self.selection_frame)
        self.from_btn.setGeometry(QtCore.QRect(20, 20, 60, 31))
        font = QtGui.QFont()
        font.setFamily("Comic Sans MS")
        font.setPointSize(10)
        self.from_btn.setFont(font)
        self.from_btn.setObjectName("from_btn")

        self.select_btn = QtWidgets.QPushButton(self.selection_frame)
        self.select_btn.setGeometry(QtCore.QRect(20, 220, 91, 50))
        font = QtGui.QFont()
        font.setFamily("Comic Sans MS")
        font.setPointSize(12)
        font.setBold(False)
        font.setItalic(False)
        font.setWeight(50)
        self.select_btn.setFont(font)
        self.select_btn.setStyleSheet("QPushButton {\n"
                                      "    font: 12pt \"Comic Sans MS\";\n"
                                      "color : white;\n"
                                      "background-color: #4CAF50;\n"
                                      "border-radius: 2px;\n"
                                      "border: 2px solid #4CAF50;\n"
                                      "}\n"
                                      "\n"
                                      "QPushButton:hover{\n"
                                      "    background-color: white;\n"
                                      "    color: black;\n"
                                      "}\n"
                                      "")
        self.select_btn.setIconSize(QtCore.QSize(100, 100))
        self.select_btn.setFlat(False)
        self.select_btn.setObjectName("select_btn")

        self.skip_btn = QtWidgets.QPushButton(self.selection_frame)
        self.skip_btn.setGeometry(QtCore.QRect(130, 220, 81, 50))
        font = QtGui.QFont()
        font.setFamily("Comic Sans MS")
        font.setPointSize(12)
        font.setBold(False)
        font.setItalic(False)
        font.setWeight(50)
        self.skip_btn.setFont(font)
        self.skip_btn.setStyleSheet("QPushButton {\n"
                                    "    font: 12pt \"Comic Sans MS\";\n"
                                    "color : white;\n"
                                    "background-color: #4CAF50;\n"
                                    "border-radius: 2px;\n"
                                    "border: 2px solid #4CAF50;\n"
                                    "}\n"
                                    "\n"
                                    "QPushButton:hover{\n"
                                    "    background-color: white;\n"
                                    "    color: black;\n"
                                    "}\n"
                                    "")
        self.skip_btn.setIconSize(QtCore.QSize(100, 100))
        self.skip_btn.setFlat(False)
        self.skip_btn.setObjectName("skip_btn")

        self.output_frame = QtWidgets.QFrame(self.centralwidget)
        self.output_frame.setGeometry(QtCore.QRect(10, 350, 261, 221))
        self.output_frame.setStyleSheet("QFrame#output_frame {\n"
                                        "    border: 3px solid #4CAF50;\n"
                                        "}")
        self.output_frame.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.output_frame.setFrameShadow(QtWidgets.QFrame.Raised)
        self.output_frame.setObjectName("output_frame")
        self.out_lbl = QtWidgets.QLabel(self.output_frame)
        self.out_lbl.setGeometry(QtCore.QRect(20, 10, 221, 81))
        font = QtGui.QFont()
        font.setFamily("Comic Sans MS")
        font.setPointSize(10)
        self.out_lbl.setFont(font)
        self.out_lbl.setStyleSheet("")
        self.out_lbl.setObjectName("out_lbl")
        self.mp3_btn = QtWidgets.QPushButton(self.output_frame)
        self.mp3_btn.setGeometry(QtCore.QRect(30, 130, 81, 50))
        font = QtGui.QFont()
        font.setFamily("Comic Sans MS")
        font.setPointSize(12)
        font.setBold(False)
        font.setItalic(False)
        font.setWeight(50)
        self.mp3_btn.setFont(font)
        self.mp3_btn.setStyleSheet("QPushButton {\n"
                                   "    font: 12pt \"Comic Sans MS\";\n"
                                   "color : white;\n"
                                   "background-color: #4CAF50;\n"
                                   "border-radius: 2px;\n"
                                   "border: 2px solid #4CAF50;\n"
                                   "}\n"
                                   "\n"
                                   "QPushButton:hover{\n"
                                   "    background-color: white;\n"
                                   "    color: black;\n"
                                   "}\n"
                                   "")
        self.mp3_btn.setIconSize(QtCore.QSize(100, 100))
        self.mp3_btn.setFlat(False)
        self.mp3_btn.setObjectName("mp3_btn")
        self.docx_btn = QtWidgets.QPushButton(self.output_frame)
        self.docx_btn.setGeometry(QtCore.QRect(150, 130, 81, 50))
        font = QtGui.QFont()
        font.setFamily("Comic Sans MS")
        font.setPointSize(12)
        font.setBold(False)
        font.setItalic(False)
        font.setWeight(50)
        self.docx_btn.setFont(font)
        self.docx_btn.setStyleSheet("QPushButton {\n"
                                    "    font: 12pt \"Comic Sans MS\";\n"
                                    "color : white;\n"
                                    "background-color: #4CAF50;\n"
                                    "border-radius: 2px;\n"
                                    "border: 2px solid #4CAF50;\n"
                                    "}\n"
                                    "\n"
                                    "QPushButton:hover{\n"
                                    "    background-color: white;\n"
                                    "    color: black;\n"
                                    "}\n"
                                    "")
        self.docx_btn.setIconSize(QtCore.QSize(100, 100))
        self.docx_btn.setFlat(False)
        self.docx_btn.setObjectName("docx_btn")

        self.fopen_frame = QtWidgets.QFrame(self.centralwidget)
        self.fopen_frame.setGeometry(QtCore.QRect(10, 350, 260, 220))
        self.fopen_frame.setStyleSheet("QFrame#fopen_frame {\n"
                                       "    border: 3px solid #4CAF50;\n"
                                       "}")
        self.fopen_frame.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.fopen_frame.setFrameShadow(QtWidgets.QFrame.Raised)
        self.fopen_frame.setObjectName("fopen_frame")
        self.out_lbl_2 = QtWidgets.QLabel(self.fopen_frame)
        self.out_lbl_2.setGeometry(QtCore.QRect(20, 10, 221, 81))
        font = QtGui.QFont()
        font.setFamily("Comic Sans MS")
        font.setPointSize(10)
        self.out_lbl_2.setFont(font)
        self.out_lbl_2.setStyleSheet("")
        self.out_lbl_2.setObjectName("out_lbl_2")

        self.open_btn = QtWidgets.QPushButton(self.fopen_frame)
        self.open_btn.setGeometry(QtCore.QRect(30, 120, 81, 50))
        font = QtGui.QFont()
        font.setFamily("Comic Sans MS")
        font.setPointSize(12)
        font.setBold(False)
        font.setItalic(False)
        font.setWeight(50)
        self.open_btn.setFont(font)
        self.open_btn.setStyleSheet("QPushButton {\n"
                                    "    font: 12pt \"Comic Sans MS\";\n"
                                    "color : white;\n"
                                    "background-color: #4CAF50;\n"
                                    "border-radius: 2px;\n"
                                    "border: 2px solid #4CAF50;\n"
                                    "}\n"
                                    "\n"
                                    "QPushButton:hover{\n"
                                    "    background-color: white;\n"
                                    "    color: black;\n"
                                    "}\n"
                                    "")
        self.open_btn.setIconSize(QtCore.QSize(100, 100))
        self.open_btn.setFlat(False)
        self.open_btn.setObjectName("open_btn")

        self.cancel_btn = QtWidgets.QPushButton(self.fopen_frame)
        self.cancel_btn.setGeometry(QtCore.QRect(150, 120, 90, 50))
        font = QtGui.QFont()
        font.setFamily("Comic Sans MS")
        font.setPointSize(12)
        font.setBold(False)
        font.setItalic(False)
        font.setWeight(50)
        self.cancel_btn.setFont(font)
        self.cancel_btn.setStyleSheet("QPushButton {\n"
                                      "    font: 12pt \"Comic Sans MS\";\n"
                                      "color : white;\n"
                                      "background-color: #4CAF50;\n"
                                      "border-radius: 2px;\n"
                                      "border: 2px solid #4CAF50;\n"
                                      "}\n"
                                      "\n"
                                      "QPushButton:hover{\n"
                                      "    background-color: white;\n"
                                      "    color: black;\n"
                                      "}\n"
                                      "")
        self.cancel_btn.setIconSize(QtCore.QSize(100, 100))
        self.cancel_btn.setFlat(False)

        self.cancel_btn.setObjectName("cancel_btn")
        UploadWindow.setCentralWidget(self.centralwidget)

        self.add_to_drop_down()

    #############################
        # self.selection_frame.setVisible(False)
        # self.output_frame.setVisible(False)
        # self.fopen_frame.setVisible(False)
        # self.fopen_frame.setVisible(False)

        self.selection_frame.close()
        self.output_frame.close()
        self.fopen_frame.close()
        self.fopen_frame.close()

        self.text_bton.clicked.connect(lambda: self.show_selection_frame(0))
        self.image_bton.clicked.connect(lambda: self.show_selection_frame(1))
        self.audio_bton.clicked.connect(lambda: self.show_selection_frame(2))

        self.select_btn.clicked.connect(
            lambda: self.open_dialogue(self.type))
        self.skip_btn.clicked.connect(
            lambda: self.skip_btn_clicked())

        self.mp3_btn.clicked.connect(self.show_fopen_frame)
        self.mp3_btn.clicked.connect(self.make_mp3)

        self.docx_btn.clicked.connect(self.show_fopen_frame)
        self.docx_btn.clicked.connect(self.make_docs)

        self.open_btn.clicked.connect(lambda: self.open_file())

        self.cancel_btn.clicked.connect(self.fopen_cancel_btn)

        self.goback_btn.clicked.connect(lambda: self.open_window(1))
        self.goback_btn.clicked.connect(UploadWindow.close)

        self.logout_btn.clicked.connect(lambda: self.open_window(2))
        self.logout_btn.clicked.connect(UploadWindow.close)
    ################################

        self.retranslateUi(UploadWindow)
        QtCore.QMetaObject.connectSlotsByName(UploadWindow)

    def retranslateUi(self, UploadWindow):
        _translate = QtCore.QCoreApplication.translate
        UploadWindow.setWindowTitle(_translate("UploadWindow", "MainWindow"))
        self.recording_label.setText(_translate(
            "UploadWindow", "<html><head/><body><p align=\"center\"><span style=\" font-size:14pt;\">Select your upload type</span></p></body></html>"))
        self.text_bton.setText(_translate("UploadWindow", "Text"))
        self.image_bton.setText(_translate("UploadWindow", "Image"))
        self.audio_bton.setText(_translate("UploadWindow", "Audio"))
        self.goback_btn.setText(_translate("UploadWindow", "Go back"))
        self.logout_btn.setText(_translate("UploadWindow", "Log Out"))

        self.to_lbl.setText(_translate(
            "UploadWindow", "<html><head/><body><p><span style=\" font-size:14pt;\">To</span></p></body></html>"))
        self.from_btn.setText(_translate(
            "UploadWindow", "<html><head/><body><p><span style=\" font-size:14pt;\">From</span></p><p><br/></p></body></html>"))
        self.select_btn.setText(_translate("UploadWindow", "Select"))
        self.skip_btn.setText(_translate("UploadWindow", "SKIP"))
        self.out_lbl.setText(_translate(
            "UploadWindow", "<html><head/><body><p><span style=\" font-size:14pt;\">Select the </span></p><p><span style=\" font-size:14pt;\">desired format</span></p></body></html>"))
        self.mp3_btn.setText(_translate("UploadWindow", "MP3"))
        self.docx_btn.setText(_translate("UploadWindow", "DOCX"))
        self.out_lbl_2.setText(_translate(
            "UploadWindow", "<html><head/><body><p><span style=\" font-size:14pt;\">File created </span></p><p><span style=\" font-size:14pt;\">successfully</span></p></body></html>"))
        self.open_btn.setText(_translate("UploadWindow", "OPEN"))
        self.cancel_btn.setText(_translate("UploadWindow", "CANCEL"))

        # default option

    def skip_btn_clicked(self):
        self.main_frame.setDisabled(False)
        self.selection_frame.close()

    def open_dialogue(self, num):
        self.from_lang = lang_dict[self.primary_lang.currentText()]
        self.to_lang = lang_dict[self.sec_lang.currentText()]

        if num == 0:
            self.open_text_file()
        elif num == 2:
            self.open_audio_file()
        else:
            self.open_image_file()

    def open_file(self):

        self.output_frame.setVisible(False)
        self.output_frame.close()

        os.system(
            f'{os.path.abspath(os.getcwd())}\\"{userdetails.directory}"\\"{self.fname}"')

    # 1 for text
    # 2 for image
    # 3 for audio

    def show_selection_frame(self, num):
        self.main_frame.setDisabled(True)
        self.type = num
        self.selection_frame.show()

    def show_fopen_frame(self):
        self.output_frame.close()

        self.fopen_frame.show()

    def fopen_cancel_btn(self):

        self.main_frame.setDisabled(False)

        self.fopen_frame.close()

    def close_window(self):
        self.lscr.close()

    def add_to_drop_down(self):

        self.sec_lang.addItems([x for x in googletrans.LANGUAGES.values()])
        self.primary_lang.addItems([x for x in googletrans.LANGUAGES.values()])
        self.primary_lang.setItemText(0, 'english')
        self.sec_lang.setItemText(0, 'hindi')
        self.primary_lang.setCurrentIndex(0)
        self.sec_lang.setCurrentIndex(0)

    # Text file upload

    def open_text_file(self):
        path = QtWidgets.QFileDialog().getOpenFileName(caption="Select a text file",
                                                       filter="Text Files (*.txt *.docx *.pdf)")[0]
        if path != "":
            self.selection_frame.setVisible(False)
            if (path.endswith('.docx')):
                text = textract.process(path, 'rb')
                text = text.decode("utf-8")
            elif (path.endswith('.txt')):
                with open(path, 'r') as f:
                    text = f.read()
            elif (path.endswith('.pdf')):
                pdf_file_obj = open(path, 'rb')
                pdf_reader = PyPDF2.PdfFileReader(pdf_file_obj)
                no_pages = pdf_reader.numPages
                text = ""
                for i in range(no_pages):
                    text += pdf_reader.getPage(i).extractText()
                pdf_file_obj.close()

            # self.selection_frame.show()
            self.output_frame.setVisible(True)
            text = translator.translate(
                text, dest=self.to_lang, src=self.from_lang).text
            self.text = text
            return

        self.text = ""

    def open_audio_file(self):

        path = QtWidgets.QFileDialog().getOpenFileName(caption="Select an audio file",
                                                       filter="Audio Files (*.wav *.mp3)")[0]
        if path != "":
            self.selection_frame.setVisible(False)

            # print(path)
            try:
                if path.endswith(".mp3"):
                    sound = pydub.AudioSegment.from_mp3(path)
                    path = path[0:len(path)-3] + "wav"
                    # print(path)
                    sound.export(path, format="wav")
                r = sr.Recognizer()
                audio_obj = sr.AudioFile(path)
                with audio_obj as source:

                    r.pause_threshold = 10
                    r.energy_threshold = 4000
                    audio = r.listen(source)

                text = r.recognize_google(audio, language=self.from_lang)
                self.text = translator.translate(
                    text, dest=self.to_lang, src=self.from_lang).text

                self.output_frame.setVisible(True)

                # self.selection_frame.show()
            except BaseException as e:
                print("The audio file isn't supported.")
                self.text = ""
                raise e

    def open_image_file(self):
        path = QtWidgets.QFileDialog().getOpenFileName(caption="Select an image file",
                                                       filter="Image Files (*.jpg *.jpeg *.jfif  *.pjpeg  *.pjp *.png)")[0]
        if (path != ""):
            self.selection_frame.setVisible(False)
            text = process_img(path, self.from_lang)
            self.output_frame.setVisible(True)
            self.text = translator.translate(
                text, dest=self.to_lang, src=self.from_lang).text

            return
        self.text = ""

    def make_docs(self):
        try:
            os.mkdir(userdetails.directory)
        except:
            pass
        self.fname = userdetails.filename + str(randint(0, 200)) + '.docx'

        document = Document()
        document.add_heading('Your translated file', 0)
        # print(self.text)
        document.add_paragraph(self.text)
        # print("inside make docs")
        # userdetails.show_user_details()
        # print(userdetails.directory)
        document.save(f"./{userdetails.directory}/{self.fname}")

    def make_mp3(self):
        try:
            os.mkdir(userdetails.directory)
        except:
            pass
        self.fname = userdetails.filename + str(randint(0, 200)) + '.mp3'
        try:
            gTTS(self.text, lang=self.to_lang).save(
                f"./{userdetails.directory}/{self.fname}")
        except:
            print("Language not supported")

        self.output_frame.setVisible(False)
